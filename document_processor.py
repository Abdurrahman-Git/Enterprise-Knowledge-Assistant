"""
WHAT THIS FILE DOES:
────────────────────
Reads different document formats (PDF, TXT, DOCX, MD) and splits them
into chunks that can be embedded and stored in the vector database.

SKILL LEARNED — Document Chunking:
────────────────────────────────────
Chunking is one of the most important decisions in RAG.

THE PROBLEM:
  You can't embed a whole 50-page document into one vector. LLMs and
  embedding models have token limits. Also, a vector for the entire doc
  would be too "averaged out" to be useful for specific questions.

THE SOLUTION — Chunking:
  Split the document into overlapping pieces:
  
  "The company was founded in 2010. We have 500 employees.
   Our main product is an HR platform. We serve 200 clients."
  
  With chunk_size=50, chunk_overlap=10:
    Chunk 1: "The company was founded in 2010. We have 500 em"
    Chunk 2: "e 500 employees. Our main product is an HR plat"
    Chunk 3: "an HR platform. We serve 200 clients."
  
  The overlap (10 chars here, 200 in real use) ensures that a sentence
  that falls at the boundary of two chunks still gets captured in at least
  one chunk — it doesn't get "split" and lose its meaning.

LANGCHAIN'S RecursiveCharacterTextSplitter:
  This is LangChain's recommended splitter. "Recursive" means it tries to
  split on paragraph breaks (\n\n) first, then newlines (\n), then spaces,
  then characters. This produces more natural chunk boundaries than just
  cutting every N characters blindly.

SKILL LEARNED — LangChain Documents:
  LangChain uses a Document object: Document(page_content="text", metadata={...})
  The metadata dict lets you store anything alongside the text:
  filename, page number, chunk index, collection name, etc.
  This metadata travels with the chunk through the entire pipeline.
"""

import os
from typing import List, Tuple
from langchain_core.documents import Document as LCDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from app.core.config import settings


class DocumentProcessor:
    """
    Reads files and splits them into LangChain Document chunks.
    
    Flow:
      Raw file bytes → extract text → split into chunks → LangChain Documents
    """

    def __init__(self):
        # RecursiveCharacterTextSplitter is LangChain's best general-purpose splitter.
        # chunk_size: max characters per chunk
        # chunk_overlap: how many characters to share between adjacent chunks
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            # These separators are tried in order — prefer splitting on paragraphs
            # before splitting mid-sentence
            separators=["\n\n", "\n", ". ", " ", ""]
        )

    def process_file(
        self,
        file_content: bytes,
        filename: str,
        collection_name: str,
        document_id: int
    ) -> Tuple[List[LCDocument], int]:
        """
        Main entry point: takes raw file bytes, returns list of LangChain Documents.
        
        Returns: (list of Document chunks, total character count)
        """
        file_ext = filename.rsplit(".", 1)[-1].lower()

        # Extract text based on file type
        if file_ext == "pdf":
            text, pages = self._extract_pdf(file_content)
        elif file_ext == "txt" or file_ext == "md":
            text = file_content.decode("utf-8", errors="ignore")
            pages = [text]
        elif file_ext == "docx":
            text, pages = self._extract_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: .{file_ext}")

        total_chars = len(text)
        
        # Split into chunks
        chunks = self._chunk_text(
            text=text,
            filename=filename,
            collection_name=collection_name,
            document_id=document_id,
            pages=pages
        )
        
        return chunks, total_chars

    def _extract_pdf(self, file_content: bytes) -> Tuple[str, List[str]]:
        """
        Extract text from PDF bytes.
        
        SKILL LEARNED — PDF Processing:
          PDFs are not text files. They're a complex format that stores
          content as positioning commands ("put character X at position Y").
          PyPDF2/pypdf reads these and reconstructs the text — but sometimes
          loses formatting, especially for scanned PDFs (which need OCR).
          
          For a portfolio project, pypdf is enough. Enterprise systems often
          use more powerful tools like AWS Textract or Azure Form Recognizer.
        """
        import io
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(file_content))
        pages = []
        full_text = ""

        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            pages.append(page_text)
            # Add page marker so we can track which page each chunk came from
            full_text += f"\n[Page {page_num + 1}]\n{page_text}"

        return full_text, pages

    def _extract_docx(self, file_content: bytes) -> Tuple[str, List[str]]:
        """Extract text from Word document."""
        import io
        from docx import Document

        doc = Document(io.BytesIO(file_content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)
        return full_text, paragraphs

    def _chunk_text(
        self,
        text: str,
        filename: str,
        collection_name: str,
        document_id: int,
        pages: List[str]
    ) -> List[LCDocument]:
        """
        Split text into overlapping chunks with metadata.
        
        SKILL LEARNED — LangChain Document Metadata:
          Every chunk carries metadata that tells us WHERE it came from.
          When we retrieve a chunk later, we can tell the user:
          "This answer came from page 3 of employee_handbook.pdf"
          
          This is crucial for enterprise use — users need to verify answers.
        """
        # Split the full document text into chunk strings
        raw_chunks = self.text_splitter.split_text(text)

        lc_documents = []
        for i, chunk_text in enumerate(raw_chunks):
            # Try to figure out which page this chunk is from
            # (simple heuristic: find page markers we inserted)
            page_num = self._estimate_page_number(chunk_text, pages)

            # LangChain Document = text content + metadata dict
            doc = LCDocument(
                page_content=chunk_text,
                metadata={
                    "filename": filename,
                    "collection_name": collection_name,
                    "document_id": str(document_id),
                    "chunk_index": i,
                    "page_number": page_num,
                    "source": filename,  # LangChain convention
                }
            )
            lc_documents.append(doc)

        return lc_documents

    def _estimate_page_number(self, chunk_text: str, pages: List[str]) -> int:
        """Estimate which page a chunk belongs to."""
        # Look for our page markers like "[Page 3]"
        import re
        match = re.search(r"\[Page (\d+)\]", chunk_text)
        if match:
            return int(match.group(1))
        return 1

    def process_text_directly(
        self,
        text: str,
        source_name: str,
        collection_name: str,
        document_id: int
    ) -> Tuple[List[LCDocument], int]:
        """Process plain text directly (not from a file upload)."""
        chunks = self._chunk_text(
            text=text,
            filename=source_name,
            collection_name=collection_name,
            document_id=document_id,
            pages=[text]
        )
        return chunks, len(text)


doc_processor = DocumentProcessor()
